"""Tests for DOCX parser using dynamically generated DOCX files."""
from io import BytesIO

from docx import Document as DocxDocument

from app.processors.docx import parse_docx


def make_test_docx() -> bytes:
    """Create a minimal DOCX with headings, paragraphs, and a table."""
    doc = DocxDocument()
    doc.core_properties.title = "Test Document"
    doc.core_properties.author = "Test Author"

    doc.add_heading("Introduction", level=1)
    doc.add_paragraph("This is the introduction paragraph with some text.")

    doc.add_heading("Background", level=2)
    doc.add_paragraph("Background information goes here.")

    doc.add_heading("Methods", level=1)
    doc.add_paragraph("Description of methods used in the study.")

    table = doc.add_table(rows=2, cols=3)
    table.rows[0].cells[0].text = "Name"
    table.rows[0].cells[1].text = "Value"
    table.rows[0].cells[2].text = "Unit"
    table.rows[1].cells[0].text = "Temperature"
    table.rows[1].cells[1].text = "25"
    table.rows[1].cells[2].text = "°C"

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def test_parse_docx_basic():
    docx_bytes = make_test_docx()
    result = parse_docx(docx_bytes)

    assert result.metadata.page_count == 1
    assert len(result.pages) == 1
    assert len(result.markdown) > 0


def test_parse_docx_headings():
    docx_bytes = make_test_docx()
    result = parse_docx(docx_bytes)

    assert len(result.headings) >= 2
    h1_texts = [h.text for h in result.headings if h.level == 1]
    assert "Introduction" in h1_texts
    assert "Methods" in h1_texts


def test_parse_docx_markdown_content():
    docx_bytes = make_test_docx()
    result = parse_docx(docx_bytes)

    assert "# Introduction" in result.markdown
    assert "## Background" in result.markdown
    assert "introduction paragraph" in result.markdown


def test_parse_docx_metadata():
    docx_bytes = make_test_docx()
    result = parse_docx(docx_bytes)

    assert result.metadata.title == "Test Document"
    assert result.metadata.author == "Test Author"


def test_parse_docx_table():
    docx_bytes = make_test_docx()
    result = parse_docx(docx_bytes)

    table_blocks = [b for b in result.pages[0].blocks if b.type == "table"]
    assert len(table_blocks) >= 1
    assert "Temperature" in table_blocks[0].content


def test_parse_docx_blocks_have_content():
    docx_bytes = make_test_docx()
    result = parse_docx(docx_bytes)

    for block in result.pages[0].blocks:
        assert block.content
        assert len(block.bbox) == 4
