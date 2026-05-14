from io import BytesIO
from typing import Optional

from docx import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.models import (
    DocumentMetadata,
    HeadingInfo,
    PageData,
    ParseResult,
    TextBlock,
)

DOCX_HEADING_STYLES = {
    "Heading 1": 1,
    "Heading 2": 2,
    "Heading 3": 3,
    "Heading 4": 4,
    "Heading 5": 5,
    "Heading 6": 6,
    "Title": 1,
    "Subtitle": 2,
}


def _paragraph_is_bold(para) -> bool:
    """Check if majority of runs in a paragraph are bold."""
    total_len = 0
    bold_len = 0
    for run in para.runs:
        text_len = len(run.text)
        total_len += text_len
        if run.bold:
            bold_len += text_len
    return bold_len > total_len * 0.5 if total_len > 0 else False


def _extract_table_markdown(table) -> str:
    """Convert a docx table to Markdown table format."""
    rows = []
    for row in table.rows:
        cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
        rows.append("| " + " | ".join(cells) + " |")

    if len(rows) >= 1:
        col_count = len(table.rows[0].cells)
        separator = "| " + " | ".join(["---"] * col_count) + " |"
        rows.insert(1, separator)

    return "\n".join(rows)


def parse_docx(file_bytes: bytes) -> ParseResult:
    """Parse a DOCX file and return structured result."""
    doc = DocxDocument(BytesIO(file_bytes))

    blocks: list[TextBlock] = []
    headings: list[HeadingInfo] = []
    md_parts: list[str] = []
    block_index = 0

    body_elements = list(doc.element.body)
    para_iter = iter(doc.paragraphs)
    table_iter = iter(doc.tables)

    current_para_idx = 0
    current_table_idx = 0

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        style_name = para.style.name if para.style else ""
        heading_level: Optional[int] = None

        for style_key, level in DOCX_HEADING_STYLES.items():
            if style_key.lower() in style_name.lower():
                heading_level = level
                break

        is_bold = _paragraph_is_bold(para)
        block_type = "heading" if heading_level else "text"

        blocks.append(TextBlock(
            type=block_type,
            content=text,
            bbox=[0, 0, 0, 0],
            heading_level=heading_level,
            is_bold=is_bold,
        ))

        if heading_level:
            headings.append(HeadingInfo(
                text=text,
                level=heading_level,
                page=1,
                block_index=block_index,
            ))
            prefix = "#" * heading_level
            md_parts.append(f"\n{prefix} {text}\n")
        else:
            md_parts.append(f"\n{text}\n")

        block_index += 1

    for table in doc.tables:
        table_md = _extract_table_markdown(table)
        if table_md:
            blocks.append(TextBlock(
                type="table",
                content=table_md,
                bbox=[0, 0, 0, 0],
            ))
            md_parts.append(f"\n{table_md}\n")
            block_index += 1

    markdown = "\n".join(md_parts).strip()

    core_props = doc.core_properties
    title = core_props.title if core_props.title else None
    author = core_props.author if core_props.author else None

    language = None
    sample = markdown[:2000]
    cjk_count = sum(1 for c in sample if "\u4e00" <= c <= "\u9fff")
    if cjk_count > len(sample) * 0.1:
        language = "zh"
    elif sample:
        language = "en"

    full_text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    page = PageData(
        page_number=1,
        width=0,
        height=0,
        text=full_text,
        blocks=blocks,
    )

    metadata = DocumentMetadata(
        title=title,
        author=author,
        page_count=1,
        language=language,
    )

    return ParseResult(
        markdown=markdown,
        pages=[page],
        headings=headings,
        metadata=metadata,
    )
